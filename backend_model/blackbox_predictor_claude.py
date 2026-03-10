import os
import json
import time
import base64
from pathlib import Path
from dotenv import load_dotenv
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

DEFAULT_MODEL_NAME = "claude-sonnet-4-20250514"
MODEL_NAME = DEFAULT_MODEL_NAME

ENV_PATH = Path(__file__).resolve().parent / ".env"
print("Loading .env from:", ENV_PATH)

load_dotenv(dotenv_path=ENV_PATH, override=True)

api_key = os.getenv("ANTHROPIC_API_KEY")

if not api_key:
    raise RuntimeError("ANTHROPIC_API_KEY not found after loading .env")

client = anthropic.Anthropic(api_key=api_key)

SCRIPT_DIR = Path(__file__).resolve().parent
TEST_DATA_DIR = SCRIPT_DIR / "test_data"
OUTPUT_DIR = SCRIPT_DIR / "case_outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# PROMPT 1: Predicts DISPOSITION + CASE DECISION
# Also accepts an optional related_cases_text string for precedent context.
# ---------------------------------------------------------------------------
PROMPT_1_TEMPLATE = """
You are serving as an appellate judge reviewing an appellate case which includes {num_docs} document(s).

CRITICAL INSTRUCTIONS:
- You must PREDICT the outcome based ONLY on the legal arguments, facts, and law presented in these documents.
- DO NOT use any external knowledge about this case, parties, or docket number.
- DO NOT simply extract or report any actual court decision that appears in the documents.
- Treat this as if YOU are the appellate court making the decision for the first time.
- Base your prediction solely on: the strength of legal arguments presented, applicable law and precedent
  discussed in the briefs, the quality of each party's legal reasoning, and the facts of the case.

The documents may include:
- Appellant's brief (arguing the lower court was wrong)
- Appellee's brief (defending the lower court decision)
- Reply briefs
- Addendums with relevant statutes, exhibits, or the lower court's decision

{related_cases_section}

Your task is to produce TWO outputs:

**OUTPUT 1: DISPOSITION**
- A single concise disposition for the case.
- Must be one of: AFFIRMED, REVERSED, VACATED, REMANDED, AFFIRMED IN PART / REVERSED IN PART, or another
  appropriate disposition if none of the above fit.

**OUTPUT 2: CASE DECISION DOCUMENT (PREDICTIVE)**
Predict and provide:
1. A written judicial opinion that decides all of the issues raised on the appeal, based on your analysis
   of the arguments presented.
2. Legal reasoning supporting your predicted decision, citing the arguments and law from the briefs.

Base your predictions on:
- Strength and persuasiveness of each party's legal arguments
- How well applicable law and precedent support each side
- Quality of legal reasoning and authority cited in the briefs
- The facts as presented and their legal significance

DO NOT base your prediction on:
- Any actual appellate decision if mentioned in the documents
- External knowledge of how this case was actually decided
- Recognition of the parties, docket numbers, or jurisdiction

Please structure your response EXACTLY as follows (do not add any text before the first delimiter):

===DISPOSITION===
[Single disposition word or phrase only — e.g. AFFIRMED]

===CASE DECISION===
[Your full written judicial opinion here]
"""

# ---------------------------------------------------------------------------
# PROMPT 2: Predicts ORAL ARGUMENT PER SIDE + CASE SUMMARY
# Injects a standard of review / complexity framework document.
# ---------------------------------------------------------------------------
PROMPT_2_TEMPLATE = """
You are serving as an appellate judge reviewing an appellate case which includes {num_docs} document(s).

CRITICAL INSTRUCTIONS:
- You must PREDICT the outcome based ONLY on the legal arguments, facts, and law presented in these documents.
- DO NOT use any external knowledge about this case, parties, or docket number.
- Treat this as if YOU are the appellate court making the decision for the first time.

The documents may include:
- Appellant's brief (arguing the lower court was wrong)
- Appellee's brief (defending the lower court decision)
- Reply briefs
- Addendums with relevant statutes, exhibits, or the lower court's decision

---
ORAL ARGUMENT STANDARD OF REVIEW — COMPLEXITY & TIME ALLOCATION GUIDELINES
{standard_of_review}
---

Your task is to produce TWO outputs:

**OUTPUT 1: ORAL ARGUMENT PER SIDE**
- Recommend a single oral argument time allocation per side based on the legal complexity of the case.
- Express as a human-readable duration string, e.g. "10 minutes", "15 minutes", "30 minutes", "1 hour".
- Base this recommendation on the complexity tier framework provided above.

**OUTPUT 2: CASE SUMMARY DOCUMENT (PREDICTIVE)**
Predict and provide:
1. A brief 3-5 sentence summary of the key legal issues on appeal.
2. Statement of the lower court's decision (what ruling is being appealed from).
3. An explanation of the case complexity and your reasoning for the oral argument time recommendation.

Please structure your response EXACTLY as follows (do not add any text before the first delimiter):

===ORAL ARGUMENT PER SIDE===
[Duration string only — e.g. 15 minutes]

===CASE SUMMARY===
[Your full predicted case summary here]
"""

STANDARD_OF_REVIEW_PLACEHOLDER = """
Refer to Standard of Review sections in the uploaded documents to determine the appropriate oral argument time allocation based on the complexity of the legal issues on appeal. Consider factors such as the number of distinct legal issues, presence of constitutional questions, factual complexity, precedent application, and any circuit splits or novel statutory interpretation required. Use your judicial reasoning to assign a complexity tier and recommend a corresponding oral argument time per side (e.g., 10 minutes, 15 minutes, 30 minutes, 1 hour) based on the guidelines provided in the placeholder document.
"""

# ---------------------------------------------------------------------------
# REDACTION PROMPT
# ---------------------------------------------------------------------------
REDACTION_PROMPT = """
Please read this legal document and return the EXACT same text, but with the following redactions:

1. Replace all docket numbers, case numbers, and case citations that identify THIS specific case
   (e.g., "No. 22-3593", "Case No. 23-1414") with "[DOCKET NUMBER REDACTED]"
2. Replace all party names (plaintiff, defendant, appellant, appellee names - both individuals and
   organizations) with generic terms:
   - Use "Appellant" for the appealing party
   - Use "Appellee" for the responding party
   - Use "Petitioner" or "Respondent" where applicable
3. Replace all judge names with "[JUDGE NAME REDACTED]"
4. Replace all attorney names and law firm names with "[ATTORNEY NAME REDACTED]" and "[LAW FIRM REDACTED]"
5. Replace specific court names (e.g., "Fifth Circuit Court of Appeals") with generic terms like
   "the Court of Appeals" or "the District Court"

IMPORTANT:
- Only redact identifying information about THIS case and its parties
- DO NOT redact case law citations, precedent names, or statutory references
  (e.g., keep "Brown v. Board of Education", "42 U.S.C. § 1983")
- Keep all legal arguments, statutes, facts, reasoning, and structure EXACTLY as written
- Preserve all legal content and formatting

Return the full redacted text with no preamble or explanation.
"""


# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def create_word_document(title, content, output_path):
    doc = Document()
    title_paragraph = doc.add_heading(title, level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphs = content.strip().split('\n')
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.style.font.size = Pt(11)
        else:
            doc.add_paragraph("")
    doc.save(output_path)
    print(f"  Created: {output_path.name}")


def create_related_cases_document(case_id, related_cases, output_path):
    doc = Document()
    title_paragraph = doc.add_heading(f"Related Cases: {case_id}", level=1)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("─" * 60)

    if not related_cases:
        p = doc.add_paragraph("No related cases available.")
        p.style.font.size = Pt(11)
    else:
        for rc in related_cases:
            docket = rc.get("docket_number", "")
            summary = rc.get("summary", "")

            docket_para = doc.add_paragraph()
            docket_run = docket_para.add_run("Docket Number: ")
            docket_run.bold = True
            docket_para.add_run(docket)
            docket_para.style.font.size = Pt(11)

            summary_label = doc.add_paragraph()
            summary_run = summary_label.add_run("Summary:")
            summary_run.bold = True
            summary_label.style.font.size = Pt(11)

            summary_para = doc.add_paragraph(summary)
            summary_para.style.font.size = Pt(11)

            doc.add_paragraph("─" * 60)

    doc.save(output_path)
    print(f"  Created: {output_path.name}")


def parse_response(response_text, delimiter_1, delimiter_2):
    """Generic two-section parser using === delimiters."""
    parts = response_text.split(delimiter_1)
    if len(parts) < 2:
        return None, None
    remainder = parts[1]
    second_parts = remainder.split(delimiter_2)
    if len(second_parts) < 2:
        return None, None
    return second_parts[0].strip(), second_parts[1].strip()


def pdf_to_text(pdf_path):
    """Extract plain text from a PDF using pdfplumber."""
    print(f"    - Extracting text from {pdf_path.name}")
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if not text.strip():
            print(f"    WARNING: No text extracted from {pdf_path.name} (may be a scanned PDF)")
            return None
        print(f"    Extracted {len(text):,} characters from {pdf_path.name}")
        return text
    except Exception as e:
        print(f"    ERROR extracting text from {pdf_path.name}: {str(e)}")
        return None


def pdf_to_base64(pdf_path):
    """Read a PDF from disk and return it as a base64-encoded string."""
    with open(pdf_path, "rb") as f:
        return base64.standard_b64encode(f.read()).decode("utf-8")


def redact_pdf_content(pdf_path, max_retries=3, conv_to_text=False):
    """Redact identifying info from a PDF via Claude."""
    if conv_to_text:
        raw_text = pdf_to_text(pdf_path)
        if not raw_text:
            return None
        print(f"    - Redacting {pdf_path.name} (text mode)")
        for attempt in range(max_retries):
            try:
                response = client.messages.create(
                    model=MODEL_NAME,
                    max_tokens=8096,
                    messages=[{
                        "role": "user",
                        "content": REDACTION_PROMPT + "\n\n" + raw_text
                    }]
                )
                return response.content[0].text
            except Exception as e:
                print(f"    Attempt {attempt + 1}/{max_retries} failed: {str(e)}")
                if attempt == max_retries - 1:
                    print(f"    Error redacting {pdf_path.name} after {max_retries} attempts")
                    return None
                wait_time = 15 * (2 ** attempt)
                print(f"    Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
    else:
        file_size_mb = pdf_path.stat().st_size / (1024 * 1024)
        print(f"    - Redacting {pdf_path.name} ({file_size_mb:.1f} MB, base64 inline mode)")
        for attempt in range(max_retries):
            try:
                pdf_data = pdf_to_base64(pdf_path)
                response = client.messages.create(
                    model=MODEL_NAME,
                    max_tokens=8096,
                    messages=[{
                        "role": "user",
                        "content": [
                            {
                                "type": "document",
                                "source": {
                                    "type": "base64",
                                    "media_type": "application/pdf",
                                    "data": pdf_data
                                }
                            },
                            {
                                "type": "text",
                                "text": REDACTION_PROMPT
                            }
                        ]
                    }]
                )
                return response.content[0].text
            except Exception as e:
                print(f"    Attempt {attempt + 1}/{max_retries} failed: {str(e)}")
                if attempt == max_retries - 1:
                    print(f"    Error redacting {pdf_path.name} after {max_retries} attempts")
                    return None
                wait_time = 15 * (2 ** attempt)
                print(f"    Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
    return None


def save_json_results(case_output_dir, case_id, disposition, oral_argument, summary, decision,
                      doc_names, call_1_success, call_2_success, related_cases):
    """Save per-case results to a JSON file."""
    results = {
        "case_id": case_id,
        "disposition": disposition,
        "oral_argument_per_side": oral_argument,
        "summary": summary,
        "decision": decision,
        "num_documents": len(doc_names),
        "documents": doc_names,
        "call_1_success": call_1_success,
        "call_2_success": call_2_success,
        "related_cases": related_cases
    }
    json_path = case_output_dir / f"{case_id}_results.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"  Created: {json_path.name}")
    return results


# ---------------------------------------------------------------------------
# CORE PROCESSING
# ---------------------------------------------------------------------------

def build_combined_content(pdf_files, conv_to_text, redact, case_output_dir, case_id):
    """Extract and combine document text from all PDFs in the case directory."""
    if redact:
        print(f"  Redacting documents...")
        redacted_texts = []
        case_output_dir.mkdir(exist_ok=True)

        for i, pdf_path in enumerate(pdf_files):
            if i > 0 and not conv_to_text:
                time.sleep(2)
            redacted_text = redact_pdf_content(pdf_path, conv_to_text=conv_to_text)
            if redacted_text:
                redacted_texts.append({"filename": pdf_path.name, "content": redacted_text})
                txt_filepath = case_output_dir / (pdf_path.stem + "_redacted.txt")
                with open(txt_filepath, 'w', encoding='utf-8') as f:
                    f.write(f"REDACTED DOCUMENT: {pdf_path.name}\n{'='*80}\n\n{redacted_text}")
                print(f"    Saved: {txt_filepath.name}")
            else:
                print(f"    WARNING: Skipping {pdf_path.name} - redaction failed")

        if not redacted_texts:
            print(f"  Failed to redact any documents for {case_id}")
            return None, []

        combined_content = ""
        for doc in redacted_texts:
            combined_content += f"\n\n{'='*60}\nDOCUMENT: {doc['filename']}\n{'='*60}\n\n{doc['content']}"
        doc_names = [doc['filename'] for doc in redacted_texts]
        return combined_content, doc_names

    else:
        if conv_to_text:
            print(f"  Extracting text from documents...")
            combined_content = ""
            doc_names = []
            for pdf_path in pdf_files:
                text = pdf_to_text(pdf_path)
                if text:
                    combined_content += f"\n\n{'='*60}\nDOCUMENT: {pdf_path.name}\n{'='*60}\n\n{text}"
                    doc_names.append(pdf_path.name)
                else:
                    print(f"  WARNING: Skipping {pdf_path.name} - text extraction failed")

            if not combined_content:
                print(f"  Failed to extract text from any documents for {case_id}")
                return None, []

            return combined_content, doc_names

        else:
            # base64 path — content embedded inline per request
            return "__FILE_UPLOAD__", [f.name for f in pdf_files]


def build_message_content(prompt, pdf_files, combined_content):
    """Assemble the messages payload for a single request."""
    if combined_content == "__FILE_UPLOAD__":
        content_blocks = [{"type": "text", "text": prompt}]
        for pdf_path in pdf_files:
            print(f"    - Encoding {pdf_path.name}")
            pdf_data = pdf_to_base64(pdf_path)
            content_blocks.append({
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": pdf_data
                }
            })
        return [{"role": "user", "content": content_blocks}]
    else:
        return [{"role": "user", "content": prompt + "\n\nDOCUMENTS:\n" + combined_content}]


def submit_and_poll_batch(requests, case_id, poll_interval=30):
    """
    Submit a mini-batch of 2 requests for a single case and poll until done.
    Returns a dict mapping custom_id -> response text (or None if that request errored).
    """
    print(f"  Submitting mini-batch ({len(requests)} requests) for case {case_id}...")
    batch = client.messages.batches.create(requests=requests)
    print(f"  Batch ID: {batch.id}")

    while True:
        batch = client.messages.batches.retrieve(batch.id)
        counts = batch.request_counts
        print(
            f"  Polling... processing={counts.processing} | "
            f"succeeded={counts.succeeded} | errored={counts.errored}"
        )
        if batch.processing_status == "ended":
            break
        time.sleep(poll_interval)

    # Collect results keyed by custom_id
    results = {}
    for result in client.messages.batches.results(batch.id):
        if result.result.type == "succeeded":
            results[result.custom_id] = result.result.message.content[0].text
        else:
            print(f"  WARNING: Request {result.custom_id} failed — {result.result.error}")
            results[result.custom_id] = None

    return results


def process_case_directory(case_dir, redact=False, conv_to_text=False):
    """
    Process a single case directory end-to-end:
      1. Extract text from all PDFs  (local, no Claude)
      2. Build both prompts
      3. Submit as a 2-request mini-batch  (bypasses 30K TPM rate limit)
      4. Poll until both results are ready
      5. Write all output files for this case immediately
    Move on to the next case only after this case is fully saved to disk.
    """
    case_id = case_dir.name
    pdf_files = sorted(case_dir.glob("*.pdf"))

    if not pdf_files:
        print(f"No PDFs found in {case_id}")
        return None

    case_output_dir = OUTPUT_DIR / case_id

    # Skip if all four output files already exist
    expected_files = [
        case_output_dir / f"{case_id}_disposition.docx",
        case_output_dir / f"{case_id}_decision.docx",
        case_output_dir / f"{case_id}_oral_argument.docx",
        case_output_dir / f"{case_id}_summary.docx",
    ]
    if case_output_dir.exists() and all(p.exists() for p in expected_files):
        print(f"\nSkipping case: {case_id} (already processed)")
        return None

    print(f"\n{'='*60}")
    print(f"Processing case: {case_id}")
    print(f"Found {len(pdf_files)} PDF(s)")
    print(f"Redaction: {'ENABLED' if redact else 'DISABLED'} | Mode: {'text conversion' if conv_to_text else 'base64 inline'}")
    print(f"{'='*60}")

    case_output_dir.mkdir(exist_ok=True)

    # ------------------------------------------------------------------
    # STEP 1: Extract text from all PDFs (local, no Claude)
    # ------------------------------------------------------------------
    combined_content, doc_names = build_combined_content(
        pdf_files, conv_to_text, redact, case_output_dir, case_id
    )
    if combined_content is None:
        return None

    num_docs = len(doc_names)
    related_cases_text = ""  # Future: populate per-case
    related_cases = []       # Future: populate per-case

    # ------------------------------------------------------------------
    # STEP 2: Build both prompts
    # ------------------------------------------------------------------
    related_cases_section = ""
    if related_cases_text.strip():
        related_cases_section = (
            "RELATED CASES FOR CONTEXT:\n"
            "The following related case summaries are provided for precedent context. "
            "Use them to inform your prediction where relevant.\n\n"
            + related_cases_text
        )

    prompt_1 = PROMPT_1_TEMPLATE.format(
        num_docs=num_docs,
        related_cases_section=related_cases_section
    )
    prompt_2 = PROMPT_2_TEMPLATE.format(
        num_docs=num_docs,
        standard_of_review=STANDARD_OF_REVIEW_PLACEHOLDER
    )

    # ------------------------------------------------------------------
    # STEP 3: Submit both prompts as a mini-batch and poll until done
    # ------------------------------------------------------------------
    batch_requests = [
        {
            "custom_id": f"{case_id}__call1",
            "params": {
                "model": MODEL_NAME,
                "max_tokens": 8096,
                "messages": build_message_content(prompt_1, pdf_files, combined_content)
            }
        },
        {
            "custom_id": f"{case_id}__call2",
            "params": {
                "model": MODEL_NAME,
                "max_tokens": 8096,
                "messages": build_message_content(prompt_2, pdf_files, combined_content)
            }
        }
    ]

    batch_results = submit_and_poll_batch(batch_requests, case_id)

    # ------------------------------------------------------------------
    # STEP 4: Parse results and write output files
    # ------------------------------------------------------------------

    # --- Call 1: Disposition + Case Decision ---
    disposition = None
    case_decision = None
    call_1_success = False

    result_1 = batch_results.get(f"{case_id}__call1")
    if result_1 is None:
        print(f"  [Call 1] ERROR: request failed, no result returned")
    else:
        disposition, case_decision = parse_response(result_1, "===DISPOSITION===", "===CASE DECISION===")
        if disposition and case_decision:
            call_1_success = True
            create_word_document(f"Disposition: {case_id}", disposition, case_output_dir / f"{case_id}_disposition.docx")
            create_word_document(f"Case Decision: {case_id}", case_decision, case_output_dir / f"{case_id}_decision.docx")
        else:
            print(f"  [Call 1] WARNING: Could not parse disposition/decision from response")
            backup_path = case_output_dir / f"{case_id}_call1_raw.txt"
            with open(backup_path, 'w', encoding='utf-8') as f:
                f.write(result_1)
            print(f"  [Call 1] Raw response saved to {backup_path.name}")

    # --- Call 2: Oral Argument + Case Summary ---
    oral_argument = None
    case_summary = None
    call_2_success = False

    result_2 = batch_results.get(f"{case_id}__call2")
    if result_2 is None:
        print(f"  [Call 2] ERROR: request failed, no result returned")
    else:
        oral_argument, case_summary = parse_response(result_2, "===ORAL ARGUMENT PER SIDE===", "===CASE SUMMARY===")
        if oral_argument and case_summary:
            call_2_success = True
            create_word_document(f"Oral Argument Per Side: {case_id}", oral_argument, case_output_dir / f"{case_id}_oral_argument.docx")
            create_word_document(f"Case Summary: {case_id}", case_summary, case_output_dir / f"{case_id}_summary.docx")
        else:
            print(f"  [Call 2] WARNING: Could not parse oral argument/summary from response")
            backup_path = case_output_dir / f"{case_id}_call2_raw.txt"
            with open(backup_path, 'w', encoding='utf-8') as f:
                f.write(result_2)
            print(f"  [Call 2] Raw response saved to {backup_path.name}")

    # --- Always write related cases doc + JSON ---
    create_related_cases_document(
        case_id=case_id,
        related_cases=related_cases,
        output_path=case_output_dir / f"{case_id}_related_cases.docx"
    )

    results = save_json_results(
        case_output_dir=case_output_dir,
        case_id=case_id,
        disposition=disposition,
        oral_argument=oral_argument,
        summary=case_summary,
        decision=case_decision,
        doc_names=doc_names,
        call_1_success=call_1_success,
        call_2_success=call_2_success,
        related_cases=related_cases
    )

    print(f"\n  Case {case_id} complete — call_1={'OK' if call_1_success else 'FAILED'} | call_2={'OK' if call_2_success else 'FAILED'}")
    return results


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main(redact=False, conv_to_text=False):
    case_dirs = sorted([d for d in TEST_DATA_DIR.iterdir() if d.is_dir()])

    print(f"Found {len(case_dirs)} case directories to process")
    print(f"Output will be saved to: {OUTPUT_DIR}")
    print(f"Redaction mode: {'ENABLED' if redact else 'DISABLED'}")
    print(f"PDF mode: {'text conversion (pdfplumber)' if conv_to_text else 'base64 inline (Anthropic)'}")

    results = []
    for case_dir in case_dirs:
        result = process_case_directory(
            case_dir,
            redact=redact,
            conv_to_text=conv_to_text
        )
        if result:
            results.append(result)

    print(f"\nProcessing complete! Processed {len(results)} cases.")
    print(f"All outputs saved to: {OUTPUT_DIR}")


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Process appellate cases with Claude.")
    parser.add_argument("--redact", action="store_true", help="Enable redaction before prediction")
    parser.add_argument("--conv_to_text", action="store_true", help="Convert PDFs to text locally with pdfplumber instead of sending as base64 to Anthropic")
    parser.add_argument("--model", type=str, default=DEFAULT_MODEL_NAME, help="Model name to use (e.g., claude-sonnet-4-20250514, claude-opus-4-20250514)")

    args = parser.parse_args()
    MODEL_NAME = args.model
    print(f"Using model: {MODEL_NAME}")

    main(redact=args.redact, conv_to_text=args.conv_to_text)